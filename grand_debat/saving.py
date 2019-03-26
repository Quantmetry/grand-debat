import numpy as np
import networkx as nx

from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
import warnings
warnings.simplefilter('ignore')


def get_theme(themes, selected_theme):
    """
    Get the sentence of the selected theme
    """
    for theme, link in themes.items():
        if int(theme[0]) == selected_theme:
            theme_sentence = theme[2:]
    return theme_sentence


def get_question(questions_id, questions_title, selected_question):
    """
    Get the sentence of the selected question
    """
    for id, title in zip(questions_id, questions_title):
        if id == selected_question:
            return title


def create_final_doc(title_sentences, titles, clean_sentences, sentences_paragraph,
                     word_embeddings, theme_sentence, question_sentence):
    """
    Get the text summarization and save results in a word document demo.docx
    """

    document = Document()

    document.add_heading('Agrégation des résultats du Grand Débat', 0)
    document.add_heading(theme_sentence, 1)
    document.add_heading(question_sentence, 2)

    for i, t in enumerate([title_sentences[x] for x in titles]):
        print(i)
        try:
            sentences_summary = [x for k, x in enumerate(clean_sentences) if sentences_paragraph.get(k, -1)==i]

            sentences_summary_emb = []
            for j in sentences_summary:
                if len(j) != 0:
                    v = sum([word_embeddings.get(w, np.zeros((100,))) for w in j.split()])/(len(j.split())+0.001)
                else:
                    v = np.zeros((100,))
                sentences_summary_emb.append(v)

            sim_mat = cosine_similarity(sentences_summary_emb)
            nx_graph = nx.from_numpy_array(sim_mat)
            try:
                scores = nx.pagerank(nx_graph)
            except:
                scores = nx.pagerank_numpy(nx_graph)
            ranked_sentences = sorted(((scores[j],s) for j, s in enumerate(sentences_summary)), reverse=True)

            sn = min(3, len(ranked_sentences))
            p_toadd = '\n'.join([ranked_sentences[j][1] for j in range(sn)])
            document.add_heading(t, level=3)
            p = document.add_paragraph(p_toadd)

        except:
            print("error for theme %i" %i)

    document.save('demo.docx')
